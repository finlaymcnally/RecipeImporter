from __future__ import annotations

import datetime as dt
import json
import logging
import re
from bisect import bisect_right
from pathlib import Path
from typing import TYPE_CHECKING, Any, Callable, Dict, List, Optional, Tuple

import yaml

try:
    import docx
except ImportError:
    docx = None

from cookimport import __version__
from cookimport.core.models import (
    ConversionReport,
    ConversionResult,
    MappingConfig,
    ParsingOverrides,
    RawArtifact,
    RecipeCandidate,
    SkippedRow,
    SourceSupport,
    WorkbookInspection,
    SheetInspection,
)
from cookimport.core.reporting import (
    ProvenanceBuilder,
    compute_file_hash,
)
from cookimport.core.source_model import normalize_source_blocks
from cookimport.parsing import cleaning, signals
from cookimport.parsing.multi_recipe_splitter import (
    MultiRecipeSplitConfig,
    split_candidate_lines,
)
from cookimport.parsing.text_section_extract import extract_sections_from_text_blob
from cookimport.plugins import registry

logger = logging.getLogger(__name__)

if TYPE_CHECKING:
    from cookimport.config.run_settings import RunSettings

# Constants for splitting
_SPLIT_DELIMITER_RE = re.compile(r"\n={3,}\s*(?:RECIPE)?\s*={3,}\n", re.IGNORECASE)
_MARKDOWN_HEADER_RE = re.compile(r"^#+\s+(.+)$", re.MULTILINE)
# Matches "1. Title" or "10. Title" at start of line
_NUMBERED_TITLE_RE = re.compile(r"^\d+\.\s+([^\n]+)$", re.MULTILINE)
_YIELD_LINE_RE = re.compile(r"^\s*(serves|yield|yields|makes)\b", re.IGNORECASE)
_INSTRUCTION_LEAD_RE = re.compile(
    r"^\s*(preheat|heat|bring|make|mix|stir|whisk|crush|cook|bake|roast|fry|grill|"
    r"blanch|season|serve|add|melt|place|put|pour|combine|fold|return|remove|drain|"
    r"peel|chop|slice|cut|toss|leave|cool|refrigerate|strain|set|beat|whip|simmer|"
    r"boil|reduce|cover|unwrap|sear|saute)\b",
    re.IGNORECASE,
)

_TABLE_HEADER_ALIASES = {
    "name": ["name", "title", "recipe", "recipe name"],
    "ingredients": ["ingredients", "ingredient", "ingredient list", "recipe/ingredients"],
    "instructions": ["instructions", "instruction", "steps", "method", "directions"],
    "description": ["description", "notes", "headnote"],
    "recipeYield": ["yield", "servings", "serves"],
    "sourceUrl": ["url", "source", "link"],
    "tags": ["tags", "keywords", "categories", "cuisine", "type", "tool"],
}

_TABLE_ALIAS_LOOKUP: dict[str, str] = {}


def _normalize_label(label: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", label.lower()).strip()


def _table_alias_lookup() -> dict[str, str]:
    if _TABLE_ALIAS_LOOKUP:
        return _TABLE_ALIAS_LOOKUP
    for field, aliases in _TABLE_HEADER_ALIASES.items():
        for alias in aliases:
            _TABLE_ALIAS_LOOKUP[_normalize_label(alias)] = field
    return _TABLE_ALIAS_LOOKUP


def _split_lines(value: Any, delimiters: list[str]) -> list[str]:
    if value is None:
        return []
    text = cleaning.normalize_text(str(value))
    for delim in delimiters:
        if delim and delim != "\n":
            text = text.replace(delim, "\n")
    lines = text.split("\n")
    return [line.strip() for line in lines if line and line.strip()]


def _normalize_ingredient_line(line: str) -> str:
    return re.sub(r"^\s*[-*]+\s+", "", line).strip()


def _normalize_instruction_line(line: str) -> str:
    cleaned = re.sub(r"^\s*[-*]+\s+", "", line)
    cleaned = re.sub(r"^\s*\d+[.)]\s+", "", cleaned)
    return cleaned.strip()


def _split_ingredients(value: Any, mapping: MappingConfig | None) -> list[str]:
    delimiters = mapping.ingredient_delimiters if mapping else ["\n"]
    lines = _split_lines(value, delimiters)
    return [_normalize_ingredient_line(line) for line in lines if _normalize_ingredient_line(line)]


def _split_instructions(value: Any, mapping: MappingConfig | None) -> list[str]:
    delimiters = mapping.instruction_delimiters if mapping else ["\n"]
    lines = _split_lines(value, delimiters)
    return [
        _normalize_instruction_line(line) for line in lines if _normalize_instruction_line(line)
    ]


def _extract_sections_from_blob(text: str) -> dict[str, list[str]]:
    return extract_sections_from_text_blob(
        text,
        ingredient_line_normalizer=_normalize_ingredient_line,
        instruction_line_normalizer=_normalize_instruction_line,
    )


def _normalize_tags(value: Any) -> list[str]:
    if value is None:
        return []
    raw = re.split(r",|\n", str(value))
    tags = [cleaning.normalize_text(tag) for tag in raw if str(tag).strip()]
    return [tag for tag in tags if tag]


def _raw_location_id(prefix: str, index: int) -> str:
    return f"{prefix}_{index}"


def _chunk_source_support(
    chunks: list[tuple[str, tuple[int, int]]],
    *,
    raw_lines: list[str],
) -> list[SourceSupport]:
    block_ids_by_line: list[str | None] = []
    next_block_index = 0
    for line in raw_lines:
        text = cleaning.normalize_text(str(line or ""))
        if not text:
            block_ids_by_line.append(None)
            continue
        block_ids_by_line.append(f"b{next_block_index}")
        next_block_index += 1

    support: list[SourceSupport] = []
    for chunk_index, (_chunk_text, line_range) in enumerate(chunks):
        start_line, end_line = line_range
        start_index = max(0, int(start_line) - 1)
        end_index = min(len(block_ids_by_line), max(start_index, int(end_line)))
        block_ids = [
            block_id
            for block_id in block_ids_by_line[start_index:end_index]
            if block_id is not None
        ]
        if not block_ids:
            continue
        support.append(
            SourceSupport(
                hintClass="proposal",
                kind="candidate_recipe_region",
                referencedBlockIds=block_ids,
                payload={
                    "chunk_index": chunk_index,
                    "start_line": int(start_line),
                    "end_line": int(end_line),
                },
                provenance={"importer": "text", "source": "multi_recipe_splitter"},
            )
        )
    return support


def _source_blocks_from_text_lines(lines: list[str]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for line_index, line in enumerate(lines):
        text = cleaning.normalize_text(str(line or ""))
        if not text:
            continue
        rows.append(
            {
                "block_id": f"b{len(rows)}",
                "order_index": len(rows),
                "text": text,
                "source_text": str(line),
                "location": {"line_index": line_index},
                "features": {"source_kind": "text_line"},
            }
        )
    return rows


def _line_start_offsets(text: str) -> list[int]:
    starts = [0]
    for index, char in enumerate(text):
        if char == "\n":
            starts.append(index + 1)
    return starts


def _trimmed_char_range(
    text: str,
    *,
    start_char: int,
    end_char: int,
) -> tuple[int, int] | None:
    start = max(0, min(len(text), int(start_char)))
    end = max(start, min(len(text), int(end_char)))
    while start < end and text[start].isspace():
        start += 1
    while end > start and text[end - 1].isspace():
        end -= 1
    if start >= end:
        return None
    return start, end


def _line_number_for_char_offset(line_starts: list[int], offset: int) -> int:
    return bisect_right(line_starts, offset) - 1


def _chunk_from_char_range(
    text: str,
    *,
    start_char: int,
    end_char: int,
    line_starts: list[int],
) -> tuple[str, tuple[int, int]] | None:
    trimmed_range = _trimmed_char_range(
        text,
        start_char=start_char,
        end_char=end_char,
    )
    if trimmed_range is None:
        return None
    trimmed_start, trimmed_end = trimmed_range
    chunk = text[trimmed_start:trimmed_end]
    start_line = _line_number_for_char_offset(line_starts, trimmed_start) + 1
    end_line = _line_number_for_char_offset(line_starts, trimmed_end - 1) + 1
    return chunk, (start_line, end_line)


def _docx_table_row_text(
    *,
    table_name: str,
    row_index: int,
    fields: dict[str, Any],
    mapping: MappingConfig | None,
) -> str:
    name = cleaning.normalize_text(str(fields.get("name") or "")).strip()
    description = cleaning.normalize_text(str(fields.get("description") or "")).strip()
    recipe_yield = cleaning.normalize_text(str(fields.get("recipeYield") or "")).strip()
    source_url = cleaning.normalize_text(str(fields.get("sourceUrl") or "")).strip()
    ingredients = _split_ingredients(fields.get("ingredients"), mapping)
    instructions = _split_instructions(fields.get("instructions"), mapping)
    tags = _normalize_tags(fields.get("tags"))

    parts: list[str] = []
    if name:
        parts.append(name)
    if description:
        parts.append(description)
    if recipe_yield:
        parts.append(f"Yield: {recipe_yield}")
    if ingredients:
        parts.append("Ingredients:")
        parts.extend(f"- {item}" for item in ingredients)
    if instructions:
        parts.append("Instructions:")
        parts.extend(f"{idx}. {item}" for idx, item in enumerate(instructions, start=1))
    if tags:
        parts.append(f"Tags: {', '.join(tags)}")
    if source_url:
        parts.append(f"Source URL: {source_url}")
    if not parts:
        parts.append(f"{table_name} row {row_index}")
    return "\n".join(parts)


def _extract_yield_phrase(line: str) -> str | None:
    match = _YIELD_LINE_RE.match(line)
    if not match:
        return None
    remainder = line[match.end():].strip(" :-")
    return remainder or line.strip()


def _is_ingredient_like(
    line: str,
    feats: dict[str, Any] | None = None,
    overrides: ParsingOverrides | None = None,
) -> bool:
    if feats is None:
        feats = signals.classify_block(line, overrides=overrides)
    if feats.get("is_ingredient_likely"):
        return True
    if feats.get("starts_with_quantity") and not feats.get("is_instruction_likely"):
        word_count = len(line.split())
        if word_count <= 5 and not feats.get("is_time"):
            return True
    if feats.get("has_unit") and not feats.get("is_instruction_likely"):
        return True
    if re.match(r"^\s*[-*•]\s+", line):
        return True
    return False


def _is_instruction_like(
    line: str,
    feats: dict[str, Any] | None = None,
    overrides: ParsingOverrides | None = None,
) -> bool:
    if feats is None:
        feats = signals.classify_block(line, overrides=overrides)
    if feats.get("is_instruction_likely"):
        return True
    if feats.get("is_ingredient_likely"):
        return False
    if _INSTRUCTION_LEAD_RE.match(line):
        return True
    word_count = len(line.split())
    if word_count >= 8 and re.search(r"[.!?]$", line.strip()):
        return True
    if word_count >= 10 and "," in line:
        return True
    return False


def _is_title_candidate(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    if _YIELD_LINE_RE.match(stripped):
        return False
    if _is_ingredient_like(stripped):
        return False
    return stripped[0].isupper() or stripped.isupper()


def _find_recipe_starts_by_yield(lines: list[str]) -> list[int]:
    starts: list[int] = []
    for idx, line in enumerate(lines):
        if not _YIELD_LINE_RE.match(line):
            continue
        prev = idx - 1
        while prev >= 0 and not lines[prev].strip():
            prev -= 1
        if prev >= 0 and _is_title_candidate(lines[prev]):
            starts.append(prev)
    if starts and starts[0] != 0:
        starts.insert(0, 0)
    return sorted(set(starts))


def _detect_docx_header_row(rows: list[list[str]]) -> int | None:
    lookup = _table_alias_lookup()
    best_idx: int | None = None
    best_score = 0
    for idx, row in enumerate(rows[:5]):
        matched_fields = {
            lookup.get(_normalize_label(cell))
            for cell in row
            if lookup.get(_normalize_label(cell))
        }
        score = len(matched_fields)
        if score > best_score:
            best_score = score
            best_idx = idx
    if best_score >= 2:
        return best_idx
    return None


def _build_docx_table_provenance(
    path: Path,
    table_name: str,
    row_index: int,
    headers: list[str],
    row_payload: dict[str, Any],
) -> dict[str, Any]:
    timestamp = dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat()
    if timestamp.endswith("+00:00"):
        timestamp = timestamp[:-6] + "Z"
    return {
        "file_path": str(path),
        "workbook": path.name,
        "sheet": table_name,
        "row_index": row_index,
        "original_headers": headers,
        "original_row": row_payload,
        "import_timestamp": timestamp,
        "converter_version": __version__,
        "extraction_method": "docx_table",
    }


class TextImporter:
    name = "text"

    def detect(self, path: Path) -> float:
        """
        Returns confidence that this is a text file we can handle.
        """
        suffix = path.suffix.lower()
        if suffix in {'.txt', '.md', '.markdown'}:
            return 0.9
        if suffix == '.docx' and docx is not None:
            return 0.95
        if suffix == '.doc':
            # Low confidence for older binary .doc input; likely to fail or be treated as text junk
            return 0.1
        return 0.0

    def inspect(self, path: Path) -> WorkbookInspection:
        """
        Analyzes the text file to determine structure (single vs multi-recipe).
        Reuses WorkbookInspection for consistency, treating the file as one 'sheet'.
        """
        try:
            if path.suffix.lower() == ".docx" and docx is not None:
                doc = docx.Document(path)
                table_rows = self._extract_docx_tables(doc)
                if table_rows:
                    recipe_rows = 0
                    tables_with_headers = 0
                    for rows in table_rows:
                        header_idx = _detect_docx_header_row(rows)
                        if header_idx is None:
                            continue
                        tables_with_headers += 1
                        recipe_rows += sum(
                            1
                            for row in rows[header_idx + 1 :]
                            if any(cell.strip() for cell in row)
                        )
                    if recipe_rows:
                        return WorkbookInspection(
                            path=str(path),
                            sheets=[
                                SheetInspection(
                                    name=path.name,
                                    layout="docx-table",
                                    headerRow=None,
                                    confidence=0.85,
                                    warnings=[
                                        f"Detected {recipe_rows} recipe row(s) across {tables_with_headers} table(s)."
                                    ],
                                )
                            ],
                            mappingStub=MappingConfig(),
                        )

            text = self._extract_text(path)
        except Exception as e:
            return WorkbookInspection(
                path=str(path),
                sheets=[],
                mappingStub=MappingConfig(),
                warnings=[f"Failed to read file: {e}"]
            )

        normalized = cleaning.normalize_text(text)
        
        candidates = self._split_recipes(normalized)
        recipe_count = len(candidates)
        
        # We'll treat the file as having one "sheet" named after the file
        sheet_name = path.name
        
        # Heuristic layout detection
        layout = "single-recipe" if recipe_count == 1 else "multi-recipe"
        
        return WorkbookInspection(
            path=str(path),
            sheets=[
                SheetInspection(
                    name=sheet_name,
                    layout=layout,
                    headerRow=None,
                    confidence=0.8,  # Arbitrary high confidence
                    warnings=[f"Detected {recipe_count} recipe candidate(s)."],
                )
            ],
            mappingStub=MappingConfig(),
        )

    def convert(
        self,
        path: Path,
        mapping: MappingConfig | None,
        progress_callback: Callable[[str], None] | None = None,
        run_settings: RunSettings | None = None,
    ) -> ConversionResult:
        report = ConversionReport()
        raw_artifacts: list[RawArtifact] = []
        overrides = mapping.parsing_overrides if mapping else None

        try:
            if path.suffix.lower() == ".docx" and docx is not None:
                if progress_callback:
                    progress_callback("Processing DOCX tables...")
                (
                    table_source_blocks,
                    table_source_support,
                    table_report,
                    table_raw,
                ) = self._convert_docx_tables(path, mapping)
                return ConversionResult(
                    recipes=[],
                    sourceBlocks=normalize_source_blocks(table_source_blocks),
                    sourceSupport=table_source_support,
                    nonRecipeBlocks=[],
                    rawArtifacts=table_raw,
                    report=table_report,
                    workbook=path.stem,
                    workbookPath=str(path),
                )

            if progress_callback:
                progress_callback("Extracting text...")
            raw_text = self._extract_text(path)
            file_hash = compute_file_hash(path)
            normalized = cleaning.normalize_text(raw_text)

            raw_artifacts.append(
                RawArtifact(
                    importer="text",
                    sourceHash=file_hash,
                    locationId="full_text",
                    extension="json",
                    content={
                        "lines": [
                            {"index": idx, "text": line}
                            for idx, line in enumerate(raw_text.splitlines())
                        ],
                        "text": raw_text,
                    },
                    metadata={"artifact_type": "extracted_text"},
                )
            )

            if progress_callback:
                progress_callback("Splitting recipes...")
            split_backend = self._resolve_multi_recipe_splitter_backend(run_settings)
            if split_backend == "off":
                total_lines = len(normalized.splitlines())
                chunks = [(normalized, (1, total_lines or 1))]
            else:
                split_result = split_candidate_lines(
                    normalized.splitlines(),
                    config=self._build_multi_recipe_split_config(
                        run_settings, backend=split_backend
                    ),
                    overrides=overrides,
                )
                chunks = self._split_by_candidate_spans(normalized, split_result.spans)
                if split_result.trace is not None:
                    raw_artifacts.append(
                        RawArtifact(
                            importer="text",
                            sourceHash=file_hash,
                            locationId="multi_recipe_split_trace",
                            extension="json",
                            content=split_result.trace,
                            metadata={
                                "artifact_type": "multi_recipe_split_trace",
                                "backend": split_backend,
                            },
                        )
                    )
            for i, (chunk_text, line_range) in enumerate(chunks):
                raw_artifacts.append(
                    RawArtifact(
                        importer="text",
                        sourceHash=file_hash,
                        locationId=_raw_location_id("chunk", i),
                        extension="json",
                        content={
                            "text": chunk_text,
                            "start_line": line_range[0],
                            "end_line": line_range[1],
                        },
                    )
                )
            source_blocks = normalize_source_blocks(
                _source_blocks_from_text_lines(raw_text.splitlines())
            )
            source_support = _chunk_source_support(
                chunks,
                raw_lines=raw_text.splitlines(),
            )
            report.total_recipes = 0
            return ConversionResult(
                recipes=[],
                sourceBlocks=source_blocks,
                sourceSupport=source_support,
                nonRecipeBlocks=[],
                rawArtifacts=raw_artifacts,
                report=report,
                workbook=path.stem,
                workbookPath=str(path),
            )

        except Exception as e:
            logger.error(f"Fatal error converting {path}: {e}")
            report.errors.append(str(e))
            return ConversionResult(
                recipes=[],
                rawArtifacts=[],
                report=report,
                workbook=path.stem,
                workbookPath=str(path),
            )

    def _extract_text(self, path: Path) -> str:
        """
        Extracts text content from the file, handling .docx and fallback for .doc.
        """
        suffix = path.suffix.lower()
        
        if suffix == '.docx':
            if docx is None:
                raise ImportError("python-docx is required for .docx files. Install it with 'pip install python-docx'.")
            try:
                doc = docx.Document(path)
                return "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                raise ValueError(f"Failed to read .docx file: {e}")
                
        elif suffix == '.doc':
            # Attempt basic text read, but likely to fail for binary
            try:
                # Try reading as text (some .doc are actually text/rtf)
                return path.read_text(encoding="utf-8", errors="replace")
            except Exception:
                # If that fails or produces binary garbage, we might need a better heuristic
                # For now, just warn/fail.
                # Check for binary signature maybe?
                raise ValueError("Binary .doc files are not fully supported. Please convert to .docx or text.")
        
        # Default text read
        return path.read_text(encoding="utf-8", errors="replace")

    def _extract_docx_tables(self, doc) -> list[list[list[str]]]:
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cleaning.normalize_text(cell.text) for cell in row.cells]
                if any(cell.strip() for cell in cells):
                    rows.append(cells)
            if rows:
                tables.append(rows)
        return tables

    def _convert_docx_tables(
        self,
        path: Path,
        mapping: MappingConfig | None,
    ) -> tuple[list[dict[str, Any]], list[SourceSupport], ConversionReport, list[RawArtifact]]:
        report = ConversionReport()
        if docx is None:
            return [], [], report, []

        try:
            doc = docx.Document(path)
        except Exception as e:
            report.errors.append(f"Failed to read .docx file: {e}")
            return [], [], report, []

        tables = self._extract_docx_tables(doc)
        if not tables:
            return [], [], report, []

        file_hash = compute_file_hash(path)
        source_blocks: list[dict[str, Any]] = []
        source_support: list[SourceSupport] = []
        raw_artifacts: list[RawArtifact] = []

        for table_idx, rows in enumerate(tables):
            header_idx = _detect_docx_header_row(rows)
            table_name = f"table_{table_idx + 1}"
            if header_idx is None:
                report.warnings.append(f"No header row detected for {table_name}.")
                continue

            headers = rows[header_idx]
            lookup = _table_alias_lookup()
            field_by_col: dict[int, str] = {}
            for idx, header in enumerate(headers):
                field = lookup.get(_normalize_label(header))
                if field:
                    field_by_col[idx] = field

            if not field_by_col:
                report.warnings.append(f"No recognizable headers found for {table_name}.")
                continue

            table_recipe_count = 0
            for row_idx, row in enumerate(rows[header_idx + 1 :], start=header_idx + 1):
                if not any(cell.strip() for cell in row):
                    continue

                row_payload: dict[str, Any] = {}
                for idx, header in enumerate(headers):
                    if idx < len(row) and header:
                        row_payload[header] = row[idx]

                fields: dict[str, Any] = {}
                all_tags: list[str] = []
                for idx, field in field_by_col.items():
                    if idx < len(row):
                        val = row[idx]
                        if field == "tags":
                            all_tags.extend(_normalize_tags(val))
                        else:
                            fields[field] = val
                if all_tags:
                    fields["tags"] = all_tags

                name = fields.get("name")
                raw_ingredients = fields.get("ingredients")
                raw_instructions = fields.get("instructions")

                if raw_ingredients and not raw_instructions:
                    sections = _extract_sections_from_blob(str(raw_ingredients))
                    if sections.get("instructions"):
                        ingredients = sections["ingredients"]
                        instructions = sections["instructions"]
                        notes_text = "\n".join(sections.get("notes", []))
                    else:
                        ingredients = _split_ingredients(raw_ingredients, mapping)
                        instructions = []
                        notes_text = ""
                else:
                    ingredients = _split_ingredients(raw_ingredients, mapping)
                    instructions = _split_instructions(raw_instructions, mapping)
                    notes_text = ""

                description = (
                    cleaning.normalize_text(str(fields.get("description")))
                    if fields.get("description")
                    else None
                )
                if notes_text:
                    if description:
                        description = f"{description}\n\n{notes_text}"
                    else:
                        description = notes_text
                recipe_yield = (
                    cleaning.normalize_text(str(fields.get("recipeYield")))
                    if fields.get("recipeYield")
                    else None
                )
                source_url = (
                    cleaning.normalize_text(str(fields.get("sourceUrl")))
                    if fields.get("sourceUrl")
                    else None
                )
                tags = _normalize_tags(fields.get("tags"))

                if not name or not str(name).strip():
                    report.skipped_rows.append(
                        SkippedRow(sheet=table_name, rowIndex=row_idx, reason="Missing name.")
                    )
                    report.missing_field_counts["name"] = (
                        report.missing_field_counts.get("name", 0) + 1
                    )
                    continue

                if not ingredients:
                    report.missing_field_counts["ingredients"] = (
                        report.missing_field_counts.get("ingredients", 0) + 1
                    )
                if not instructions:
                    report.missing_field_counts["instructions"] = (
                        report.missing_field_counts.get("instructions", 0) + 1
                    )

                block_id = f"b{len(source_blocks)}"
                source_blocks.append(
                    {
                        "block_id": block_id,
                        "order_index": len(source_blocks),
                        "text": _docx_table_row_text(
                            table_name=table_name,
                            row_index=row_idx,
                            fields={
                                "name": name,
                                "description": description or "",
                                "recipeYield": recipe_yield or "",
                                "sourceUrl": source_url or "",
                                "tags": tags,
                                "ingredients": "\n".join(ingredients),
                                "instructions": "\n".join(instructions),
                            },
                            mapping=mapping,
                        ),
                        "location": {"table_id": table_name, "table_row_index": row_idx},
                        "features": {"source_kind": "docx_table_row"},
                        "provenance": {"importer": "text", "source_hash": file_hash},
                    }
                )
                source_support.append(
                    SourceSupport(
                        hintClass="evidence",
                        kind="docx_table_row",
                        referencedBlockIds=[block_id],
                        payload={
                            "table": table_name,
                            "row_index": row_idx,
                            "headers": headers,
                            "row": row_payload,
                        },
                        provenance={"importer": "text", "source": "docx_table"},
                    )
                )
                raw_artifacts.append(
                    RawArtifact(
                        importer="text",
                        sourceHash=file_hash,
                        locationId=f"{table_name}_r{row_idx}",
                        extension="json",
                        content={
                            "table": table_name,
                            "row_index": row_idx,
                            "headers": headers,
                            "row": row_payload,
                        },
                    )
                )
                table_recipe_count += 1

            report.per_sheet_counts[table_name] = report.per_sheet_counts.get(
                table_name, 0
            ) + table_recipe_count

        if raw_artifacts:
            raw_artifacts.append(
                RawArtifact(
                    importer="text",
                    sourceHash=file_hash,
                    locationId="full_rows",
                    extension="json",
                    content={
                        "rows": [
                            artifact.content
                            for artifact in raw_artifacts
                            if isinstance(artifact.content, dict)
                            and "row" in artifact.content
                        ]
                    },
                    metadata={"artifact_type": "extracted_rows"},
                )
            )

        report.total_recipes = 0
        return source_blocks, source_support, report, raw_artifacts

    def _split_recipes(self, text: str) -> List[Tuple[str, Tuple[int, int]]]:
        """
        Splits text into recipe chunks.
        Returns list of (text_chunk, (start_line, end_line)).
        """
        lines = text.splitlines()
        yield_starts = _find_recipe_starts_by_yield(lines)
        if len(yield_starts) > 1:
            return self._split_by_line_indices(text, yield_starts)

        # Strategy 1: Explicit delimiters
        if _SPLIT_DELIMITER_RE.search(text):
            return self._split_by_regex(text, _SPLIT_DELIMITER_RE)
            
        # Strategy 2: Markdown Headers (if multiple H1s are present)
        h1_matches = list(_MARKDOWN_HEADER_RE.finditer(text))
        if len(h1_matches) > 1:
            # Check if they look like recipe titles (not too long)
            # This is a simplification; we might assume top-level headers are recipes
            return self._split_by_positions(text, [m.start() for m in h1_matches])

        # Strategy 3: Numbered List of Titles (e.g. "1. Recipe Name")
        # Heuristic: If we see multiple numbered items, and NO ingredients headers,
        # it's likely a list of recipes/ideas.
        numbered_matches = list(_NUMBERED_TITLE_RE.finditer(text))
        has_ingredients = re.search(r"^Ingredients\b", text, re.IGNORECASE | re.MULTILINE)
        
        if len(numbered_matches) > 1 and not has_ingredients:
            return self._split_by_positions(text, [m.start() for m in numbered_matches])
            
        # Strategy 4: Default to single recipe
        return [(text, (1, len(text.splitlines())))]

    def _resolve_multi_recipe_splitter_backend(
        self, run_settings: RunSettings | None
    ) -> str:
        raw_backend = getattr(getattr(run_settings, "multi_recipe_splitter", None), "value", None)
        if raw_backend is None and run_settings is not None:
            raw_backend = getattr(run_settings, "multi_recipe_splitter", None)
        normalized = str(raw_backend or "rules_v1").strip().lower().replace("-", "_")
        if normalized in {"off", "rules_v1"}:
            return normalized
        return "rules_v1"

    def _build_multi_recipe_split_config(
        self,
        run_settings: RunSettings | None,
        *,
        backend: str,
    ) -> MultiRecipeSplitConfig:
        min_ingredient_lines = getattr(run_settings, "multi_recipe_min_ingredient_lines", 1)
        min_instruction_lines = getattr(
            run_settings, "multi_recipe_min_instruction_lines", 1
        )
        for_the_guardrail = getattr(
            run_settings, "multi_recipe_for_the_guardrail", True
        )
        trace = getattr(run_settings, "multi_recipe_trace", False)
        return MultiRecipeSplitConfig(
            backend=backend,
            min_ingredient_lines=max(0, int(min_ingredient_lines or 0)),
            min_instruction_lines=max(0, int(min_instruction_lines or 0)),
            enable_for_the_guardrail=bool(for_the_guardrail),
            trace=bool(trace),
        )

    def _split_by_candidate_spans(
        self,
        text: str,
        spans: tuple[Any, ...],
    ) -> List[Tuple[str, Tuple[int, int]]]:
        lines = text.splitlines(keepends=True)
        if not lines:
            return []

        line_starts: list[int] = []
        pos = 0
        for line in lines:
            line_starts.append(pos)
            pos += len(line)

        chunks: list[tuple[str, tuple[int, int]]] = []
        total_lines = len(lines)
        for span in spans:
            start_line = max(0, int(getattr(span, "start", 0)))
            end_line = min(total_lines, int(getattr(span, "end", total_lines)))
            if end_line <= start_line:
                continue
            start_char = line_starts[start_line]
            end_char = line_starts[end_line] if end_line < total_lines else len(text)
            chunk = text[start_char:end_char].strip()
            if not chunk:
                continue
            chunks.append((chunk, (start_line + 1, end_line)))

        if chunks:
            return chunks
        return [(text, (1, total_lines))]

    def _split_by_regex(self, text: str, pattern: re.Pattern) -> List[Tuple[str, Tuple[int, int]]]:
        chunks = []
        line_starts = _line_start_offsets(text)
        matches = list(pattern.finditer(text))
        if not matches:
            total_lines = max(1, len(text.splitlines()))
            return [(text, (1, total_lines))]

        current_start = 0
        for match in matches:
            chunk = _chunk_from_char_range(
                text,
                start_char=current_start,
                end_char=match.start(),
                line_starts=line_starts,
            )
            if chunk is not None:
                chunks.append(chunk)
            current_start = match.end()

        last_chunk = _chunk_from_char_range(
            text,
            start_char=current_start,
            end_char=len(text),
            line_starts=line_starts,
        )
        if last_chunk is not None:
            chunks.append(last_chunk)

        return chunks

    def _split_by_positions(self, text: str, positions: List[int]) -> List[Tuple[str, Tuple[int, int]]]:
        chunks = []
        line_starts = _line_start_offsets(text)
        cleaned_positions = sorted(
            {
                max(0, min(len(text), int(position)))
                for position in positions
            }
        )
        for i in range(len(cleaned_positions)):
            start = cleaned_positions[i]
            end = cleaned_positions[i + 1] if i + 1 < len(cleaned_positions) else len(text)
            chunk = _chunk_from_char_range(
                text,
                start_char=start,
                end_char=end,
                line_starts=line_starts,
            )
            if chunk is not None:
                chunks.append(chunk)
        return chunks

    def _split_by_line_indices(
        self,
        text: str,
        start_lines: list[int],
    ) -> List[Tuple[str, Tuple[int, int]]]:
        lines = text.splitlines(keepends=True)
        if not lines:
            return []
        line_starts: list[int] = []
        pos = 0
        for line in lines:
            line_starts.append(pos)
            pos += len(line)

        chunks: list[tuple[str, tuple[int, int]]] = []
        starts = sorted(set(start_lines))
        for idx, start_line in enumerate(starts):
            if start_line < 0 or start_line >= len(lines):
                continue
            end_line = starts[idx + 1] if idx + 1 < len(starts) else len(lines)
            start_char = line_starts[start_line]
            end_char = line_starts[end_line] if end_line < len(lines) else len(text)
            chunk = text[start_char:end_char].strip()
            if chunk:
                chunks.append((chunk, (start_line + 1, end_line)))
        return chunks

    def _parse_chunk(
        self,
        text: str,
        overrides: ParsingOverrides | None = None,
    ) -> RecipeCandidate:
        """
        Parses a single recipe chunk into a candidate.
        """
        lines = text.splitlines()
        
        # 1. Frontmatter
        frontmatter = {}
        content_lines = lines
        if text.startswith("---"):
            try:
                # Find end of frontmatter
                end_idx = -1
                for i in range(1, len(lines)):
                    if lines[i].strip() == "---":
                        end_idx = i
                        break
                if end_idx > 0:
                    fm_text = "\n".join(lines[1:end_idx])
                    frontmatter = yaml.safe_load(fm_text) or {}
                    content_lines = lines[end_idx+1:]
            except Exception as e:
                logger.warning(f"Failed to parse frontmatter: {e}")

        # 2. Identify Sections
        name = frontmatter.get("title")
        ingredients: List[str] = []
        instructions: List[str] = []
        description_lines: List[str] = []
        tags = frontmatter.get("tags", [])
        if isinstance(tags, str):
            tags = [t.strip() for t in tags.split(",")]

        current_section = "description" # Default start
        recipe_yield = (
            str(frontmatter.get("servings") or frontmatter.get("yield") or "")
            if frontmatter.get("servings") or frontmatter.get("yield")
            else None
        )
        
        # Heuristic: First line is title if not in frontmatter
        start_idx = 0
        if not name:
            for i, line in enumerate(content_lines):
                stripped = line.strip()
                if stripped:
                    # Cleanup title candidates
                    name = re.sub(r"^#+\s*", "", stripped) # Remove markdown headers
                    name = re.sub(r"^\d+\.\s*", "", name)  # Remove leading numbering "1. "
                    name = re.sub(r"^Title:\s*", "", name, flags=re.IGNORECASE) # Remove "Title:" prefix
                    start_idx = i + 1
                    break

        remaining_lines = [line.strip() for line in content_lines[start_idx:] if line.strip()]
        if remaining_lines:
            has_headers = False
            cleaned_lines: list[str] = []
            for line in remaining_lines:
                block_feats = signals.classify_block(line, overrides=overrides)
                if block_feats["is_ingredient_header"] or block_feats["is_instruction_header"]:
                    has_headers = True
                if recipe_yield is None:
                    extracted = _extract_yield_phrase(line)
                    if extracted is not None:
                        recipe_yield = extracted
                        continue
                cleaned_lines.append(line)

            if has_headers:
                for line in cleaned_lines:
                    block_feats = signals.classify_block(line, overrides=overrides)

                    if block_feats["is_ingredient_header"]:
                        current_section = "ingredients"
                        continue
                    if block_feats["is_instruction_header"]:
                        current_section = "instructions"
                        continue
                    if block_feats["is_header_likely"] and "notes" in line.lower():
                        current_section = "description"
                        continue

                    if current_section == "ingredients":
                        clean_line = re.sub(r"^\s*[-*•]\s*", "", line.strip())
                        ingredients.append(clean_line)
                    elif current_section == "instructions":
                        clean_line = re.sub(r"^\s*(?:\d+[.)]|\*|-)\s*", "", line.strip())
                        instructions.append(clean_line)
                    else:
                        description_lines.append(line.strip())
            else:
                feats_list = [
                    signals.classify_block(line, overrides=overrides)
                    for line in cleaned_lines
                ]
                ingredient_indices = [
                    i for i, line in enumerate(cleaned_lines)
                    if _is_ingredient_like(line, feats_list[i], overrides=overrides)
                ]
                instruction_indices = [
                    i for i, line in enumerate(cleaned_lines)
                    if _is_instruction_like(line, feats_list[i], overrides=overrides)
                ]

                ingredient_start = ingredient_indices[0] if ingredient_indices else None
                instruction_start = None
                if ingredient_start is not None:
                    for idx in instruction_indices:
                        if idx > ingredient_start:
                            instruction_start = idx
                            break

                if ingredient_start is not None:
                    description_lines.extend(cleaned_lines[:ingredient_start])
                    ingredient_block = cleaned_lines[
                        ingredient_start:instruction_start or len(cleaned_lines)
                    ]
                    instruction_block = (
                        cleaned_lines[instruction_start:] if instruction_start is not None else []
                    )
                elif instruction_indices:
                    instruction_start = instruction_indices[0]
                    description_lines.extend(cleaned_lines[:instruction_start])
                    ingredient_block = []
                    instruction_block = cleaned_lines[instruction_start:]
                else:
                    description_lines.extend(cleaned_lines)
                    ingredient_block = []
                    instruction_block = []

                for idx, line in enumerate(ingredient_block):
                    feats = signals.classify_block(line, overrides=overrides)
                    if _is_instruction_like(line, feats, overrides=overrides) and not _is_ingredient_like(line, feats, overrides=overrides):
                        instruction_block = ingredient_block[idx:] + instruction_block
                        ingredient_block = ingredient_block[:idx]
                        break

                for line in ingredient_block:
                    clean_line = re.sub(r"^\s*[-*•]\s*", "", line.strip())
                    if clean_line:
                        ingredients.append(clean_line)

                for line in instruction_block:
                    clean_line = re.sub(r"^\s*(?:\d+[.)]|\*|-)\s*", "", line.strip())
                    if clean_line:
                        instructions.append(clean_line)

        return RecipeCandidate(
            name=name or "Untitled Recipe",
            ingredients=ingredients,
            instructions=instructions,
            description="\n".join(description_lines) if description_lines else None,
            recipeYield=recipe_yield or "",
            sourceUrl=frontmatter.get("source"),
            tags=tags
        )

registry.register(TextImporter())
